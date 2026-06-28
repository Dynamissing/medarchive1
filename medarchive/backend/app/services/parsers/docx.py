from __future__ import annotations

import shutil
import subprocess
import tempfile
import zipfile
from dataclasses import dataclass
from pathlib import Path
from typing import Any
from xml.etree import ElementTree

from docx import Document

from app.schemas.parsed_document import ParsedDocumentResult
from app.services.parsers.base import DocumentParser, ParserInput
from app.services.parsers.spreadsheet import find_libreoffice

TRACKED_CHANGE_TAGS = {
    "ins",
    "del",
    "moveFrom",
    "moveTo",
    "cellIns",
    "cellDel",
    "cellMerge",
    "tblPrChange",
    "trPrChange",
    "tcPrChange",
    "pPrChange",
    "rPrChange",
}


class DocxParser(DocumentParser):
    parser_name = "docx"
    parser_format = "docx"

    def parse(self, parser_input: ParserInput) -> ParsedDocumentResult:
        extraction = extract_docx(parser_input.source_path)
        warnings = list(extraction.warnings)

        fallback_path: Path | None = None
        if extraction.has_tracked_changes or is_poor_quality(extraction):
            try:
                fallback_path = convert_docx_to_docx(parser_input.source_path)
            except RuntimeError as exc:
                warnings.append(str(exc))
            else:
                fallback_extraction = extract_docx(fallback_path)
                if is_better_quality(fallback_extraction, extraction):
                    extraction = fallback_extraction.with_warnings(warnings)
                    warnings = list(extraction.warnings)
                warnings.append("LibreOffice fallback DOCX extraction was attempted.")
            finally:
                if fallback_path is not None:
                    shutil.rmtree(fallback_path.parent, ignore_errors=True)

        return ParsedDocumentResult(
            parser_name=self.parser_name,
            parser_format="docx",
            status="parsed",
            source_file_asset_id=parser_input.file_asset_id,
            source_path=str(parser_input.source_path),
            extracted_text=extraction.raw_text,
            tables=extraction.tables,
            metadata={
                "paragraph_count": len(extraction.paragraphs),
                "table_count": len(extraction.tables),
                "has_tracked_changes": extraction.has_tracked_changes,
                "fallback_recommended": extraction.has_tracked_changes or is_poor_quality(extraction),
                "mime_type": parser_input.mime_type,
                "extension": parser_input.extension,
            },
            warnings=warnings,
        )


@dataclass(frozen=True)
class DocxExtraction:
    paragraphs: list[dict[str, Any]]
    tables: list[dict[str, Any]]
    raw_text: str
    has_tracked_changes: bool
    warnings: list[str]

    def with_warnings(self, warnings: list[str]) -> DocxExtraction:
        return DocxExtraction(
            paragraphs=self.paragraphs,
            tables=self.tables,
            raw_text=self.raw_text,
            has_tracked_changes=self.has_tracked_changes,
            warnings=warnings,
        )


def extract_docx(path: Path) -> DocxExtraction:
    document = Document(path)
    paragraphs: list[dict[str, Any]] = []
    raw_parts: list[str] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        paragraphs.append(
            {
                "paragraph_index": paragraph_index,
                "text": text,
                "locator": f"paragraph:{paragraph_index}",
            }
        )
        raw_parts.append(text)

    tables: list[dict[str, Any]] = []
    for table_index, table in enumerate(document.tables):
        rows: list[dict[str, Any]] = []
        for row_index, row in enumerate(table.rows):
            values = [cell.text.strip() for cell in row.cells]
            if all(value == "" for value in values):
                continue
            rows.append(
                {
                    "row_index": row_index,
                    "values": values,
                    "locator": f"table:{table_index}:row:{row_index}",
                }
            )
            raw_parts.append("\t".join(values))
        tables.append(
            {
                "table_index": table_index,
                "row_count": len(rows),
                "rows": rows,
                "locator": f"table:{table_index}",
            }
        )

    warnings: list[str] = []
    has_tracked_changes = detect_tracked_changes(path)
    if has_tracked_changes:
        warnings.append("Tracked changes detected in DOCX OOXML; LibreOffice fallback is recommended.")

    return DocxExtraction(
        paragraphs=paragraphs,
        tables=tables,
        raw_text="\n".join(raw_parts),
        has_tracked_changes=has_tracked_changes,
        warnings=warnings,
    )


def detect_tracked_changes(path: Path) -> bool:
    try:
        with zipfile.ZipFile(path) as archive:
            xml_names = [name for name in archive.namelist() if name.startswith("word/") and name.endswith(".xml")]
            for xml_name in xml_names:
                with archive.open(xml_name) as xml_file:
                    for _event, element in ElementTree.iterparse(xml_file, events=("start",)):
                        tag_name = element.tag.rsplit("}", 1)[-1]
                        if tag_name in TRACKED_CHANGE_TAGS:
                            return True
    except (OSError, zipfile.BadZipFile, ElementTree.ParseError):
        return False
    return False


def is_poor_quality(extraction: DocxExtraction) -> bool:
    return not extraction.raw_text.strip() and not any(table["row_count"] for table in extraction.tables)


def is_better_quality(candidate: DocxExtraction, current: DocxExtraction) -> bool:
    candidate_score = len(candidate.raw_text.strip()) + sum(table["row_count"] for table in candidate.tables) * 25
    current_score = len(current.raw_text.strip()) + sum(table["row_count"] for table in current.tables) * 25
    return candidate_score > current_score


def convert_docx_to_docx(source_path: Path) -> Path:
    executable = find_libreoffice()
    if executable is None:
        raise RuntimeError("LibreOffice executable was not found for DOCX fallback conversion.")

    output_dir = Path(tempfile.mkdtemp(prefix="medarchive-docx-"))
    command = [
        str(executable),
        "--headless",
        "--convert-to",
        "docx",
        "--outdir",
        str(output_dir),
        str(source_path),
    ]
    completed = subprocess.run(command, capture_output=True, text=True, check=False)
    if completed.returncode != 0:
        shutil.rmtree(output_dir, ignore_errors=True)
        raise RuntimeError(f"LibreOffice DOCX fallback conversion failed: {completed.stderr or completed.stdout}")

    converted = output_dir / source_path.name
    if not converted.exists():
        matches = list(output_dir.glob("*.docx"))
        if not matches:
            shutil.rmtree(output_dir, ignore_errors=True)
            raise RuntimeError("LibreOffice DOCX fallback did not produce a DOCX file.")
        converted = matches[0]
    return converted
