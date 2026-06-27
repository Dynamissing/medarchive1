from __future__ import annotations

import shutil
import zipfile
from pathlib import Path

from docx import Document

from app.services.parsers.base import ParserInput
from app.services.parsers.docx import DocxParser, detect_tracked_changes


def create_docx_with_table(path: Path) -> None:
    document = Document()
    document.add_paragraph("Clinic 1 price list")
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Code"
    table.cell(0, 1).text = "Service"
    table.cell(0, 2).text = "Price"
    table.cell(1, 0).text = "A1"
    table.cell(1, 1).text = "Consultation"
    table.cell(1, 2).text = "500"
    table.cell(2, 0).text = "B1"
    table.cell(2, 1).text = "MRI"
    table.cell(2, 2).text = "1500"
    document.add_paragraph("Footer note")
    document.save(path)


def add_tracked_change_marker(source_path: Path, target_path: Path) -> None:
    shutil.copyfile(source_path, target_path)
    with zipfile.ZipFile(target_path, "a") as archive:
        archive.writestr(
            "word/commentsExtended.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:ins xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
                'w:id="1" w:author="tester" />'
            ),
        )


def test_docx_parser_extracts_paragraphs_tables_raw_text_and_locators(tmp_path: Path) -> None:
    docx_path = tmp_path / "clinic.docx"
    create_docx_with_table(docx_path)

    result = DocxParser().parse(ParserInput(parser_format="docx", source_path=docx_path))

    assert result.status == "parsed"
    assert result.parser_format == "docx"
    assert "Clinic 1 price list" in result.extracted_text
    assert "Consultation" in result.extracted_text
    assert result.metadata["paragraph_count"] == 2
    assert result.metadata["table_count"] == 1
    assert result.metadata["has_tracked_changes"] is False
    assert result.tables[0]["table_index"] == 0
    assert result.tables[0]["rows"][1]["row_index"] == 1
    assert result.tables[0]["rows"][1]["locator"] == "table:0:row:1"
    assert result.tables[0]["rows"][1]["values"] == ["A1", "Consultation", "500"]


def test_docx_tracked_change_detection_and_fallback_warning(tmp_path: Path) -> None:
    clean_path = tmp_path / "clean.docx"
    tracked_path = tmp_path / "tracked.docx"
    create_docx_with_table(clean_path)
    add_tracked_change_marker(clean_path, tracked_path)

    assert detect_tracked_changes(tracked_path) is True

    result = DocxParser().parse(ParserInput(parser_format="docx", source_path=tracked_path))

    assert result.metadata["has_tracked_changes"] is True
    assert result.metadata["fallback_recommended"] is True
    assert "Tracked changes detected in DOCX OOXML; LibreOffice fallback is recommended." in result.warnings
    assert any("LibreOffice executable was not found" in warning for warning in result.warnings)
