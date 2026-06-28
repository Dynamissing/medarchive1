from __future__ import annotations

import zipfile
from pathlib import Path

from docx import Document
from openpyxl import Workbook
from PIL import Image
from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas


def create_synthetic_xlsx(path: Path) -> Path:
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Prices"
    sheet.append(["Code", "Service", "Resident price", "Nonresident price"])
    sheet.append(["A-1", "Blood test", 1000, 1200])
    sheet.append(["B-1", "MRI brain", 5000, 6000])
    workbook.save(path)
    return path


def create_synthetic_docx(path: Path) -> Path:
    document = Document()
    document.add_paragraph("Synthetic clinic price list")
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).text = "Code"
    table.cell(0, 1).text = "Service"
    table.cell(0, 2).text = "Price"
    table.cell(1, 0).text = "A-1"
    table.cell(1, 1).text = "Blood test"
    table.cell(1, 2).text = "1000"
    table.cell(2, 0).text = "B-1"
    table.cell(2, 1).text = "MRI brain"
    table.cell(2, 2).text = "5000"
    document.save(path)
    return path


def create_synthetic_text_pdf(path: Path) -> Path:
    pdf = canvas.Canvas(str(path), pagesize=letter)
    pdf.drawString(72, 740, "Synthetic clinic price list")
    pdf.drawString(72, 710, "001 Blood test 1000")
    pdf.drawString(72, 690, "002 MRI brain 5000")
    pdf.save()
    return path


def create_synthetic_scanned_pdf(path: Path) -> Path:
    image_path = path.with_suffix(".png")
    image = Image.new("RGB", (320, 140), "white")
    image.save(image_path)
    pdf = canvas.Canvas(str(path), pagesize=letter)
    pdf.drawImage(str(image_path), 72, 600, width=320, height=140)
    pdf.save()
    return path


def create_archive_with_xlsx(path: Path, workbook_path: Path) -> Path:
    with zipfile.ZipFile(path, "w") as archive:
        archive.write(workbook_path, arcname="clinic/prices.xlsx")
    return path


def fake_ocr_data(confidence: int = 90) -> dict[str, list]:
    words = ["001", "Blood", "test", "1000", "002", "MRI", "brain", "5000"]
    return {
        "text": words,
        "conf": [confidence] * 4 + [55] * 4,
        "left": [10, 50, 100, 180, 10, 50, 100, 180],
        "block_num": [1] * 8,
        "par_num": [1] * 8,
        "line_num": [1, 1, 1, 1, 2, 2, 2, 2],
    }
